try:
    __import__("pysqlite3")
    import sys

    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    # This happens on your Windows machine, which is totally fine!
    pass

import streamlit as st
import os
import re
from groq import Groq
from dotenv import load_dotenv
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import chromadb
from chromadb.utils import embedding_functions
from docx import Document
from io import BytesIO

# --- 1. SETUP ---
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# USE A PERSISTENT PATH COMPATIBLE WITH CLOUD
DB_PATH = "./my_vector_db"
chroma_client = chromadb.PersistentClient(path=DB_PATH)
sentence_transformer_ef = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name="all-MiniLM-L6-v2"
)

# --- 2. SESSION STATE MANAGEMENT ---
if "messages" not in st.session_state:
    st.session_state.messages = []
if "indexed_files" not in st.session_state:
    st.session_state.indexed_files = []
if "processing" not in st.session_state:
    st.session_state.processing = False


# --- 3. EXPORT LOGIC ---
def generate_docx(messages):
    doc = Document()
    doc.add_heading("Enterprise AI Knowledge Report", 0)
    for msg in messages:
        role_name = "USER" if msg["role"] == "user" else "AI ASSISTANT"
        doc.add_heading(role_name, level=2)
        clean_text = (
            msg["content"].replace("**", "").replace("#", "").replace("|", "  ")
        )
        clean_text = re.sub(r" +", " ", clean_text)
        doc.add_paragraph(clean_text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- 4. TEXT PROCESSING ---
def get_pdf_text(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text


def get_text_chunks(raw_text):
    return RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200
    ).split_text(raw_text)


# --- 5. UI ---
st.set_page_config(page_title="AI Knowledge Hub PRO", layout="wide")
st.title("🏛️ Enterprise Knowledge Hub")

with st.sidebar:
    if st.session_state.processing:
        if st.button("Emergency Unlock UI"):
            st.session_state.processing = False
            st.rerun()

    persona = st.selectbox(
        "Choose AI Role:",
        ["General Assistant", "Legal Expert", "Financial Advisor", "IT Engineer"],
    )
    st.divider()

    uploaded_files = st.file_uploader(
        "Upload PDFs",
        type="pdf",
        accept_multiple_files=True,
        key="file_uploader_mobile",
    )

    requires_indexing = uploaded_files and len(uploaded_files) != len(
        st.session_state.indexed_files
    )

    # --- INDEXING ACTION ---
    if st.button(
        "Index All Files", disabled=st.session_state.processing or not uploaded_files
    ):
        st.session_state.processing = True
        try:
            # Delete and recreate to ensure a fresh index on every "Index All" click
            try:
                chroma_client.delete_collection("library")
            except:
                pass

            collection = chroma_client.create_collection(
                name="library", embedding_function=sentence_transformer_ef
            )

            st.session_state.indexed_files = []

            for file in uploaded_files:
                with st.spinner(f"Indexing {file.name}..."):
                    text = get_pdf_text(file)
                    chunks = get_text_chunks(text)
                    collection.add(
                        documents=chunks,
                        ids=[f"{file.name}_{i}" for i in range(len(chunks))],
                        metadatas=[{"source": file.name} for _ in range(len(chunks))],
                    )
                    st.session_state.indexed_files.append(file.name)
            st.success("Library Fully Synced!")
        except Exception as e:
            st.error(f"Sync Error: {e}")
        finally:
            st.session_state.processing = False
            st.rerun()

    if st.session_state.indexed_files:
        st.write("✅ **Documents in Memory:**")
        for f in st.session_state.indexed_files:
            st.write(f"- {f}")

    if st.button(
        "⚖️ Compare Documents", disabled=len(st.session_state.indexed_files) < 2
    ):
        st.session_state.processing = True
        try:
            collection = chroma_client.get_collection(
                name="library", embedding_function=sentence_transformer_ef
            )
            all_context = ""
            for fname in st.session_state.indexed_files:
                data = collection.get(where={"source": fname}, limit=5)
                all_context += f"\nFILE: {fname}\n{data['documents']}\n"

            res = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "user",
                        "content": f"Compare these files in a detailed report:\n{all_context}",
                    }
                ],
            )
            st.session_state.messages.append(
                {
                    "role": "assistant",
                    "content": f"### Comparison Report\n{res.choices[0].message.content}",
                }
            )
        except Exception as e:
            st.error(f"Comparison Error: {e}")
        finally:
            st.session_state.processing = False
            st.rerun()

    if st.button("🗑️ Clear Everything"):
        try:
            chroma_client.delete_collection("library")
            st.session_state.messages, st.session_state.indexed_files = [], []
            st.rerun()
        except:
            st.rerun()

# --- 6. CHAT LOGIC ---
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

if requires_indexing:
    st.warning("Action Required: Please click 'Index All Files'.")
    chat_disabled = True
elif not st.session_state.indexed_files:
    st.info("Upload and index a PDF to start.")
    chat_disabled = True
else:
    chat_disabled = False

if prompt := st.chat_input(
    "Query your knowledge hub...", disabled=st.session_state.processing or chat_disabled
):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    try:
        collection = chroma_client.get_collection(
            name="library", embedding_function=sentence_transformer_ef
        )
        results = collection.query(query_texts=[prompt], n_results=5)
        context = "\n\n".join(
            [
                f"[{m['source']}]: {d}"
                for d, m in zip(results["documents"][0], results["metadatas"][0])
            ]
        )

        res = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "user",
                    "content": f"You are a {persona}. Context:\n{context}\n\nQuestion: {prompt}",
                }
            ],
        )
        ai_ans = res.choices[0].message.content
        st.session_state.messages.append({"role": "assistant", "content": ai_ans})
        st.rerun()
    except Exception as e:
        st.error(f"System Error: {e}")
